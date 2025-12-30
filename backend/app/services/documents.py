"""Document text extraction and caching service.

Handles extracting text from various document formats and caching the
results to avoid reprocessing. Supports TXT, MD, CSV, JSON, PDF, DOCX,
and XLSX files.
"""

from __future__ import annotations

import csv
import json
import re
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from .storage import StorageManager


class DocumentProcessor:
    """Extracts and caches text content from documents.
    
    Supports multiple file formats and maintains a text cache in the
    .doc_matrix folder to avoid reprocessing unchanged files.
    
    Attributes:
        storage: Storage manager for cache operations.
    """
    
    def __init__(self, storage: StorageManager) -> None:
        """Initialize document processor with storage manager.
        
        Args:
            storage: Storage manager instance.
        """
        self.storage = storage
    
    def get_document_text(
        self, 
        filename: str, 
        force_refresh: bool = False
    ) -> Dict[str, Any]:
        """Get extracted text for a document, using cache when available.
        
        Args:
            filename: Name of the document file.
            force_refresh: If True, reprocess even if cached.
            
        Returns:
            Dictionary with text content, pages, and metadata.
        """
        doc_path = self.storage.root / filename
        if not doc_path.exists():
            raise FileNotFoundError(f"Document not found: {filename}")
        
        cache_path = self.storage.get_cache_path(filename)
        current_mtime = self.storage.get_file_mtime(doc_path)
        
        # Check cache validity
        if not force_refresh and cache_path.exists():
            cached = self.storage.read_json(cache_path)
            if cached and cached.get("mtime") == current_mtime:
                return cached
        
        # Extract text based on file type
        ext = doc_path.suffix.lower()
        extractor = self._get_extractor(ext)
        
        if extractor is None:
            raise ValueError(f"Unsupported file type: {ext}")
        
        result = extractor(doc_path)
        result["mtime"] = current_mtime
        result["filename"] = filename
        
        # Cache the result
        self.storage.write_json(cache_path, result)
        
        return result
    
    def _get_extractor(self, extension: str):
        """Get the appropriate extractor function for a file extension.
        
        Args:
            extension: File extension (including dot).
            
        Returns:
            Extractor function or None if unsupported.
        """
        extractors = {
            ".txt": self._extract_text,
            ".md": self._extract_text,
            ".csv": self._extract_csv,
            ".json": self._extract_json,
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".xlsx": self._extract_xlsx,
        }
        return extractors.get(extension)
    
    def _extract_text(self, path: Path) -> Dict[str, Any]:
        """Extract text from plain text files.
        
        Args:
            path: Path to the text file.
            
        Returns:
            Extraction result with text and metadata.
        """
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
        
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "char_start": 0, 
                       "char_end": len(text)}],
            "char_count": len(text),
        }
    
    def _extract_csv(self, path: Path) -> Dict[str, Any]:
        """Extract text from CSV files.
        
        Args:
            path: Path to the CSV file.
            
        Returns:
            Extraction result with text representation.
        """
        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = path.read_text(encoding="latin-1")
        
        # Parse CSV and convert to readable text
        reader = csv.reader(StringIO(content))
        rows = list(reader)
        
        if not rows:
            return {"text": "", "pages": [], "char_count": 0}
        
        # Format as readable text with headers
        headers = rows[0] if rows else []
        text_lines = []
        
        for i, row in enumerate(rows):
            if i == 0:
                text_lines.append("Headers: " + ", ".join(row))
            else:
                row_text = "; ".join(
                    f"{headers[j] if j < len(headers) else f'Col{j}'}: {val}"
                    for j, val in enumerate(row)
                )
                text_lines.append(f"Row {i}: {row_text}")
        
        text = "\n".join(text_lines)
        
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "char_start": 0,
                       "char_end": len(text)}],
            "char_count": len(text),
            "row_count": len(rows),
        }
    
    def _extract_json(self, path: Path) -> Dict[str, Any]:
        """Extract text from JSON files.
        
        Args:
            path: Path to the JSON file.
            
        Returns:
            Extraction result with pretty-printed JSON.
        """
        content = path.read_text(encoding="utf-8")
        
        try:
            data = json.loads(content)
            text = json.dumps(data, indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            text = content
        
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "char_start": 0,
                       "char_end": len(text)}],
            "char_count": len(text),
        }
    
    def _extract_pdf(self, path: Path) -> Dict[str, Any]:
        """Extract text from PDF files.
        
        Args:
            path: Path to the PDF file.
            
        Returns:
            Extraction result with text per page.
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF extraction")
        
        reader = PdfReader(str(path))
        pages = []
        full_text = []
        char_offset = 0
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            page_text = self._clean_text(page_text)
            
            pages.append({
                "page": i + 1,
                "text": page_text,
                "char_start": char_offset,
                "char_end": char_offset + len(page_text),
            })
            
            full_text.append(page_text)
            char_offset += len(page_text) + 1  # +1 for newline
        
        combined_text = "\n".join(full_text)
        
        return {
            "text": combined_text,
            "pages": pages,
            "char_count": len(combined_text),
            "page_count": len(pages),
        }
    
    def _extract_docx(self, path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files.
        
        Args:
            path: Path to the DOCX file.
            
        Returns:
            Extraction result with text content.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX extraction")
        
        doc = Document(str(path))
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        
        text = "\n\n".join(paragraphs)
        
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "char_start": 0,
                       "char_end": len(text)}],
            "char_count": len(text),
            "paragraph_count": len(paragraphs),
        }
    
    def _extract_xlsx(self, path: Path) -> Dict[str, Any]:
        """Extract text from XLSX files.
        
        Args:
            path: Path to the XLSX file.
            
        Returns:
            Extraction result with text from all sheets.
        """
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl is required for XLSX extraction")
        
        wb = load_workbook(str(path), read_only=True, data_only=True)
        sheets_text = []
        pages = []
        char_offset = 0
        
        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            rows_text = [f"=== Sheet: {sheet_name} ==="]
            
            for row in sheet.iter_rows(values_only=True):
                # Filter out None values and convert to strings
                row_values = [
                    str(cell) if cell is not None else "" 
                    for cell in row
                ]
                if any(v.strip() for v in row_values):
                    rows_text.append(" | ".join(row_values))
            
            sheet_text = "\n".join(rows_text)
            sheets_text.append(sheet_text)
            
            pages.append({
                "page": sheet_idx + 1,
                "sheet_name": sheet_name,
                "text": sheet_text,
                "char_start": char_offset,
                "char_end": char_offset + len(sheet_text),
            })
            
            char_offset += len(sheet_text) + 2  # +2 for double newline
        
        wb.close()
        
        text = "\n\n".join(sheets_text)
        
        return {
            "text": text,
            "pages": pages,
            "char_count": len(text),
            "sheet_count": len(pages),
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text by normalizing whitespace.
        
        Args:
            text: Raw extracted text.
            
        Returns:
            Cleaned text with normalized whitespace.
        """
        # Replace multiple spaces with single space
        text = re.sub(r" +", " ", text)
        # Replace multiple newlines with double newline
        text = re.sub(r"\n\s*\n", "\n\n", text)
        return text.strip()
    
    def find_text_location(
        self, 
        filename: str, 
        search_text: str
    ) -> Optional[Dict[str, Any]]:
        """Find the location of text within a document.
        
        Args:
            filename: Name of the document.
            search_text: Text to search for.
            
        Returns:
            Location info with page and character positions, or None.
        """
        doc_data = self.get_document_text(filename)
        full_text = doc_data.get("text", "")
        
        # Find the text in the full document
        start_pos = full_text.find(search_text)
        if start_pos == -1:
            # Try case-insensitive search
            start_pos = full_text.lower().find(search_text.lower())
        
        if start_pos == -1:
            return None
        
        end_pos = start_pos + len(search_text)
        
        # Find which page contains this text
        page_num = 1
        for page in doc_data.get("pages", []):
            if page["char_start"] <= start_pos < page["char_end"]:
                page_num = page.get("page", page.get("sheet_name", 1))
                break
        
        return {
            "filename": filename,
            "page": page_num,
            "char_start": start_pos,
            "char_end": end_pos,
            "text": search_text,
        }
    
    def get_context_around(
        self, 
        filename: str, 
        char_start: int, 
        char_end: int,
        context_chars: int = 100
    ) -> str:
        """Get surrounding context for a text location.
        
        Args:
            filename: Name of the document.
            char_start: Start character position.
            char_end: End character position.
            context_chars: Number of context characters on each side.
            
        Returns:
            Text with surrounding context.
        """
        doc_data = self.get_document_text(filename)
        full_text = doc_data.get("text", "")
        
        context_start = max(0, char_start - context_chars)
        context_end = min(len(full_text), char_end + context_chars)
        
        prefix = "..." if context_start > 0 else ""
        suffix = "..." if context_end < len(full_text) else ""
        
        return prefix + full_text[context_start:context_end] + suffix

